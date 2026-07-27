import argparse
import csv
import re
import sys
from collections import Counter
from pathlib import Path

try:
    import docx  # python-docx
except ImportError:
    docx = None

DEFAULT_DICT_PATH = Path(__file__).parent / "dictionary_v3.csv"
TOKEN_RE = re.compile(r"[A-Za-z][A-Za-z'-]*")
SUPPORTED_EXTS = {".txt", ".docx"}


def load_dictionary(path):
    categories = {}
    with open(path, newline="", encoding="utf-8-sig") as f:
        for row in csv.reader(f):
            if not row or not row[0].strip():
                continue 
            name = row[0].strip()
            patterns = [w.strip().lower() for w in row[1:] if w.strip()]
            exact = {p for p in patterns if not p.endswith("*")}
            prefixes = [p[:-1] for p in patterns if p.endswith("*")]
            categories[name] = {"exact": exact, "prefixes": prefixes}
    return categories


def read_text(path):
    """Read text from a .txt or .docx file."""
    path = Path(path)
    ext = path.suffix.lower()
    if ext == ".docx":
        if docx is None:
            raise RuntimeError(
                "python-docx is not installed. Run: pip install python-docx"
            )
        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    return path.read_text(encoding="utf-8", errors="ignore")


def tokenize(text):
    return [t.lower() for t in TOKEN_RE.findall(text)]


def count_categories(text, categories):
    tokens = tokenize(text)
    results = {}
    for name, spec in categories.items():
        matched = Counter()
        for tok in tokens:
            if tok in spec["exact"]:
                matched[tok] += 1
                continue
            for prefix in spec["prefixes"]:
                if tok.startswith(prefix):
                    matched[tok] += 1
                    break
        results[name] = {"count": sum(matched.values()), "words": matched}
    return results, len(tokens)


def print_report(results, total_words):
    name_width = max(len(n) for n in results) + 2
    print(f"Total words in letter: {total_words}\n")
    print(f"{'Category':<{name_width}}{'Count':>7}   Matched words (with frequency)")
    print("-" * 70)
    for name, data in results.items():
        top = ", ".join(f"{w}×{c}" if c > 1 else w for w, c in data["words"].most_common())
        print(f"{name:<{name_width}}{data['count']:>7}   {top}")


def find_input_files(folder, recursive=False):
    folder = Path(folder)
    pattern = "**/*" if recursive else "*"
    files = [
        p for p in sorted(folder.glob(pattern))
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXTS
    ]
    return files


def process_folder(folder, categories, out_path, recursive=False):
    files = find_input_files(folder, recursive=recursive)
    if not files:
        print(f"No .txt or .docx files found in {folder}", file=sys.stderr)
        return

    category_names = list(categories.keys())
    fieldnames = ["filename", "total_words"]
    for name in category_names:
        fieldnames.append(f"{name}_count")
        fieldnames.append(f"{name}_words")

    rows = []
    for f in files:
        try:
            text = read_text(f)
        except Exception as e:
            print(f"Skipping {f.name}: {e}", file=sys.stderr)
            continue

        results, total_words = count_categories(text, categories)
        row = {"filename": f.name, "total_words": total_words}
        for name in category_names:
            data = results[name]
            row[f"{name}_count"] = data["count"]
            row[f"{name}_words"] = ", ".join(
                f"{w}×{c}" if c > 1 else w for w, c in data["words"].most_common()
            )
        rows.append(row)
        print(f"Processed {f.name}")

    with open(out_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)

    print(f"\nDone. Wrote {len(rows)} rows to {out_path}")


def main():
    parser = argparse.ArgumentParser(
        description="Count dictionary category matches in one letter, or batch-process a folder of .txt/.docx letters into a combined CSV."
    )
    parser.add_argument(
        "input",
        nargs="?",
        help="Path to a single letter file, or a folder to batch-process. Omit to read stdin.",
    )
    parser.add_argument("--dict", default=str(DEFAULT_DICT_PATH), help="Path to dictionary CSV")
    parser.add_argument("--out", default="results.csv", help="Output CSV path (batch/folder mode only)")
    parser.add_argument("--recursive", action="store_true", help="Include subfolders when batch-processing a folder")
    args = parser.parse_args()

    categories = load_dictionary(args.dict)

    if args.input and Path(args.input).is_dir():
        process_folder(args.input, categories, args.out, recursive=args.recursive)
        return

    if args.input:
        text = read_text(args.input)
    else:
        text = sys.stdin.read()

    results, total_words = count_categories(text, categories)
    print_report(results, total_words)


if __name__ == "__main__":
    main()